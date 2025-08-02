from os.path import exists
import re
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from upload_file import upload_file
from pathlib import Path
import io

def load_templates():
    """Loads presentation templates"""
    custom_template = Path(__file__).parent / "template.docx"
    if exists(custom_template):
        template = custom_template
    else:
        # If no template found, create a blank document
        return None
    return str(template)

def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    """Adds a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def parse_inline_formatting(text, paragraph):
    """Parse inline markdown formatting like **bold**, *italic*, and [links](url)"""
    # First handle escape characters
    text = handle_escapes(text)

    # Handle line breaks (two spaces at end of line)
    # Split by line breaks while preserving them
    line_parts = text.split('  \n')

    for line_idx, line_part in enumerate(line_parts):
        if not line_part and line_idx == len(line_parts) - 1:
            continue

        # Split text by formatting markers while preserving the markers
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', line_part)

        for part in parts:
            if not part:
                continue

            # Bold text (**text**)
            if part.startswith('**') and part.endswith('**'):
                bold_text = part[2:-2]
                paragraph.add_run(bold_text).bold = True
            # Italic text (*text*)
            elif part.startswith('*') and part.endswith('*'):
                italic_text = part[1:-1]
                paragraph.add_run(italic_text).italic = True
            # Inline code (`code`)
            elif part.startswith('`') and part.endswith('`'):
                code_text = part[1:-1]
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
            # Links [text](url)
            elif part.startswith('[') and '](' in part and part.endswith(')'):
                link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if link_match:
                    link_text, url = link_match.groups()
                    add_hyperlink(paragraph, link_text, url)
            else:
                # Plain text
                paragraph.add_run(part)

        # Add line break if this isn't the last part
        if line_idx < len(line_parts) - 1:
            paragraph.add_run().add_break()

def handle_escapes(text):
    """Handle backslash escaped characters"""
    # Replace escaped characters with placeholders first to avoid conflicts
    escape_map = {}
    placeholder_counter = 0

    def replace_escape(match):
        nonlocal placeholder_counter
        escaped_char = match.group(1)
        placeholder = f"__ESC_{placeholder_counter}__"
        escape_map[placeholder] = escaped_char
        placeholder_counter += 1
        return placeholder

    # Find and replace all escaped characters
    text = re.sub(r'\\(.)', replace_escape, text)

    # After all other processing, restore the escaped characters
    for placeholder, char in escape_map.items():
        text = text.replace(placeholder, char)

    return text

def parse_table(lines, start_idx):
    """Parse markdown table and return the table data and next line index"""
    table_lines = []
    i = start_idx

    # Find all table lines
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
        else:
            break

    if len(table_lines) < 2:  # Need at least header and separator
        return None, start_idx + 1

    # Parse table data
    table_data = []
    for line in table_lines:
        # Skip separator line (contains dashes)
        if '---' in line or ':-:' in line or ':--' in line or '--:' in line:
            continue

        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
        table_data.append(cells)

    return table_data, i

def add_table_to_doc(table_data, doc):
    """Add table data to Word document"""
    if not table_data:
        return

    rows = len(table_data)
    cols = max(len(row) for row in table_data) if table_data else 0

    word_table = doc.add_table(rows=rows, cols=cols)
    word_table.style = 'Table Grid'

    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = word_table.cell(i, j)
                if cell.paragraphs:
                    cell.paragraphs[0].clear()

                cell_paragraph = cell.paragraphs[0]
                parse_inline_formatting(cell_text, cell_paragraph)

def markdown_to_word(markdown_content):
    """Convert Markdown to Word document."""
    path = load_templates()

    # Create document with or without template
    if path:
        doc = Document(path)
    else:
        doc = Document()  # Create blank document if no template
        print("Warning: No template found, creating blank document")

    # Split content into lines, but preserve line breaks within paragraphs
    lines = markdown_content.split('\n')
    i = 0

    try:
        while i < len(lines):
            line = lines[i]

            # Handle multiple empty lines (preserve spacing)
            if not line.strip():
                empty_line_count = 0
                start_empty = i

                # Count consecutive empty lines
                while i < len(lines) and not lines[i].strip():
                    empty_line_count += 1
                    i += 1

                # Add appropriate spacing based on number of empty lines
                if empty_line_count == 1:
                    # Single empty line = normal paragraph break (already handled by next iteration)
                    pass
                elif empty_line_count >= 2:
                    # Multiple empty lines = add extra spacing
                    # Add one empty paragraph for each additional empty line beyond the first
                    for _ in range(empty_line_count - 1):
                        doc.add_paragraph()

                continue

            # Check if this line ends with two spaces (line break)
            if line.endswith('  '):
                # Collect lines that are part of the same paragraph (connected by line breaks)
                paragraph_lines = []
                while i < len(lines):
                    current_line = lines[i]
                    if not current_line.strip():  # Empty line ends the paragraph
                        break

                    paragraph_lines.append(current_line)
                    i += 1

                    # If line doesn't end with two spaces, this paragraph is complete
                    if not current_line.endswith('  '):
                        break

                # Join lines with line break markers
                full_text = '  \n'.join(paragraph_lines)

                # Determine what type of content this is
                first_line = paragraph_lines[0].strip()

                # Headers
                if first_line.startswith('#'):
                    header_level = len(first_line) - len(first_line.lstrip('#'))
                    header_text = first_line.lstrip('#').strip()
                    heading = doc.add_heading('', level=min(header_level, 6))
                    parse_inline_formatting(header_text, heading)

                # Block quotes
                elif first_line.startswith('>'):
                    quote_text = full_text[1:].strip()  # Remove > from beginning
                    quote_paragraph = doc.add_paragraph()
                    quote_paragraph.style = 'Quote'
                    parse_inline_formatting(quote_text, quote_paragraph)

                # Regular paragraph with line breaks
                else:
                    paragraph = doc.add_paragraph()
                    parse_inline_formatting(full_text, paragraph)

                continue

            line = line.strip()

            # Headers
            if line.startswith('#'):
                header_level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                heading = doc.add_heading('', level=min(header_level, 6))
                parse_inline_formatting(header_text, heading)
                i += 1

            # Tables
            elif line.startswith('|'):
                table_data, i = parse_table(lines, i)
                if table_data:
                    add_table_to_doc(table_data, doc)

            # Ordered lists
            elif re.match(r'^\d+\.\s+', line):
                i = process_list_items(lines, i, doc, True, 0)

            # Unordered lists
            elif re.match(r'^[-*+]\s+', line):
                i = process_list_items(lines, i, doc, False, 0)

            # Horizontal rule
            elif line.startswith('---') or line.startswith('***'):
                # Add a horizontal line (simplified as empty paragraph with border)
                paragraph = doc.add_paragraph()
                i += 1

            # Block quotes (useful for legal citations)
            elif line.startswith('>'):
                quote_text = line[1:].strip()
                quote_paragraph = doc.add_paragraph()
                quote_paragraph.style = 'Quote'
                parse_inline_formatting(quote_text, quote_paragraph)
                i += 1

            # Regular paragraphs
            else:
                paragraph = doc.add_paragraph()
                parse_inline_formatting(line, paragraph)
                i += 1

    except Exception as e:
        print(f"Error in parsing markdown: {e}")
        import traceback
        traceback.print_exc()
        return f"Error in parsing markdown: {e}"

    # Save the document to BytesIO and upload
    try:
        # Save to BytesIO object
        file_object = io.BytesIO()
        doc.save(file_object)
        file_object.seek(0)

        # Upload and get result
        result = upload_file(file_object, "docx")
        file_object.close()

        print(f"Word document uploaded successfully")
        return result
    except Exception as e:
        print(f"Error saving/uploading Word document: {e}")
        import traceback
        traceback.print_exc()
        return f"Error saving/uploading Word document: {e}"

def process_list_items(lines, start_idx, doc, is_ordered=False, level=0):
    """Process markdown list items with proper Word numbering"""
    bullet_styles = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
    number_styles = ['List Number', 'List Number 2', 'List Number 3']

    style_array = number_styles if is_ordered else bullet_styles
    style = style_array[min(level, len(style_array) - 1)]

    i = start_idx

    while i < len(lines):
        line = lines[i].strip()

        # Determine indentation level from original line
        original_line = lines[i]
        indent = len(original_line) - len(original_line.lstrip())
        current_level = indent // 3  # Use 3 spaces per level to match typical markdown indentation

        # If indentation doesn't match our expected level, this item doesn't belong to this list
        if current_level != level:
            break

        # Check if this is a list item at our current level
        if is_ordered:
            list_match = re.match(r'^\d+\.\s+(.+)', line)
        else:
            list_match = re.match(r'^[-*+]\s+(.+)', line)

        if not list_match:
            break

        item_text = list_match.group(1)

        # Use Word's built-in list formatting - it handles numbering restart automatically
        paragraph = doc.add_paragraph(style=style)
        parse_inline_formatting(item_text, paragraph)

        i += 1

        # Look ahead for nested items
        while i < len(lines):
            if i >= len(lines):
                break

            next_line = lines[i].strip()
            if not next_line:
                i += 1
                continue

            next_original = lines[i]
            next_indent = len(next_original) - len(next_original.lstrip())
            next_level = next_indent // 3  # Use 3 spaces per level

            if next_level > level:
                # This is a nested item - process the nested list
                if re.match(r'^\d+\.\s+', next_line):
                    i = process_list_items(lines, i, doc, True, next_level)
                elif re.match(r'^[-*+]\s+', next_line):
                    i = process_list_items(lines, i, doc, False, next_level)
                else:
                    # Not a list item, stop processing nested items
                    break
            elif next_level == level:
                # Same level item - continue in this loop
                break
            else:
                # Lower level - return to parent
                break

    return i
